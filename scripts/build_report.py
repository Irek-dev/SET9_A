#!/usr/bin/env python3
from __future__ import annotations

import csv
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
AVERAGED = ROOT / "data" / "averaged_results.csv"
RAW = ROOT / "data" / "raw_results.csv"
FIGURES = ROOT / "figures"
REPORT = ROOT / "report" / "report_A1.docx"
IDS_FILE = ROOT / "submissions_ids.txt"
REPOSITORY_URL = "https://github.com/Irek-dev/SET9_A.git"

SEED = 20260524
ALPHABET = "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789!@#%:;^&*()-"
ARRAY_LABELS = {
    "random": "случайный",
    "reverse_sorted": "обратно отсортированный",
    "nearly_sorted": "почти отсортированный",
}
ALGORITHM_LABELS = {
    "QuickSort": "Обычный Quick Sort",
    "MergeSort": "Обычный Merge Sort",
    "StringQuickSort": "String Quick Sort",
    "StringMergeSort": "String Merge Sort",
    "MSDRadixSort": "MSD Radix Sort",
    "MSDRadixQuickSort": "MSD Radix + Quick",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def style_table(table, widths: list[int]) -> None:
    table.autofit = False
    table.allow_autofit = False
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            if row_index == 0:
                set_cell_shading(cell, "FFFFFF")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)


def load_rows() -> list[dict[str, str]]:
    with AVERAGED.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def load_submission_ids() -> dict[str, str]:
    ids = {"A1m": "НЕ ЗАПОЛНЕНО", "A1q": "НЕ ЗАПОЛНЕНО", "A1r": "НЕ ЗАПОЛНЕНО", "A1rq": "НЕ ЗАПОЛНЕНО"}
    if IDS_FILE.exists():
        for line in IDS_FILE.read_text(encoding="utf-8").splitlines():
            if ":" in line:
                key, value = line.split(":", 1)
                key = key.strip()
                if key in ids:
                    ids[key] = value.strip() or "НЕ ЗАПОЛНЕНО"
    return ids


def format_num(value: float, digits: int = 1) -> str:
    return f"{value:,.{digits}f}".replace(",", " ")


def analyze(rows: list[dict[str, str]]) -> dict[str, object]:
    by_type_size: dict[tuple[str, int], list[dict[str, str]]] = defaultdict(list)
    by_type: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        by_type_size[(row["array_type"], int(row["size"]))].append(row)
        by_type[row["array_type"]].append(row)

    summary = {}
    for array_type in ARRAY_LABELS:
        at_3000 = by_type_size[(array_type, 3000)]
        fastest_3000 = min(at_3000, key=lambda row: float(row["mean_time_us"]))
        least_chars_3000 = min(at_3000, key=lambda row: float(row["mean_char_comparisons"]))
        slowest_3000 = max(at_3000, key=lambda row: float(row["mean_time_us"]))
        summary[array_type] = {
            "fastest_3000": fastest_3000,
            "least_chars_3000": least_chars_3000,
            "slowest_3000": slowest_3000,
            "speedup_3000": float(slowest_3000["mean_time_us"]) / max(float(fastest_3000["mean_time_us"]), 1.0),
        }

    algorithm_means = {}
    for algorithm in ALGORITHM_LABELS:
        values = [row for row in rows if row["algorithm"] == algorithm]
        algorithm_means[algorithm] = {
            "time": sum(float(row["mean_time_us"]) for row in values) / len(values),
            "chars": sum(float(row["mean_char_comparisons"]) for row in values) / len(values),
        }

    return {"summary": summary, "algorithm_means": algorithm_means}


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Title", 20, "000000", 0, 8),
        ("Heading 1", 15, "000000", 14, 6),
        ("Heading 2", 12, "000000", 10, 4),
        ("Heading 3", 11, "000000", 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_ids_table(doc: Document, ids: dict[str, str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Задача"
    table.rows[0].cells[1].text = "ID успешной посылки Codeforces"
    for key in ["A1m", "A1q", "A1r", "A1rq"]:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = ids[key]
    style_table(table, [1800, 7400])


def add_summary_table(doc: Document, analysis: dict[str, object]) -> None:
    summary = analysis["summary"]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Тип массива", "Быстрее всех при n=3000", "Время, мс", "Минимум символов", "Симв., тыс."]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for array_type in ARRAY_LABELS:
        item = summary[array_type]
        fastest = item["fastest_3000"]
        least_chars = item["least_chars_3000"]
        cells = table.add_row().cells
        cells[0].text = ARRAY_LABELS[array_type]
        cells[1].text = ALGORITHM_LABELS[fastest["algorithm"]]
        cells[2].text = format_num(float(fastest["mean_time_us"]) / 1000.0, 2)
        cells[3].text = ALGORITHM_LABELS[least_chars["algorithm"]]
        cells[4].text = format_num(float(least_chars["mean_char_comparisons"]) / 1000.0, 1)
    style_table(table, [1700, 2550, 1400, 2550, 1000])


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_black_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    new_run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    new_run.append(properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_graphs(doc: Document) -> None:
    for array_type in ARRAY_LABELS:
        add_black_heading(doc, f"Графики: {ARRAY_LABELS[array_type]} массив", level=2)
        for prefix, caption in [("time", "среднее время работы"), ("comparisons", "символьные сравнения")]:
            path = FIGURES / f"{prefix}_{array_type}.png"
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.25))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p = doc.add_paragraph(f"Рисунок: {caption}, {ARRAY_LABELS[array_type]} массив.")
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption_p.runs:
                    run.italic = True


def build_report() -> None:
    rows = load_rows()
    analysis = analyze(rows)
    ids = load_submission_ids()

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Normal"]
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("Отчёт A1: анализ строковых сортировок")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    add_black_heading(doc, "Методика", level=1)
    add_paragraph(
        doc,
        f"StringGenerator: алфавит из {len(ALPHABET)} символов "
        "A..Z, a..z, 0..9, «!@#%:;^&*()-». Длина строк: 10..200."
    )
    add_paragraph(
        doc,
        "Размеры массивов: 100..3000 с шагом 100. Для каждого типа сначала генерировался массив "
        "из 3000 строк, затем брались префиксы нужной длины. Типы: случайный, обратно "
        "отсортированный, почти отсортированный."
    )
    add_paragraph(
        doc,
        f"StringSortTester: 10 запусков на каждую комбинацию, seed {SEED}. Измерялись время "
        "и число посимвольных сравнений. Для Radix Sort считались обращения к текущему символу."
    )
    add_paragraph(
        doc,
        "Обычный Quick Sort и String Quick Sort используют одинаковый pivot: середину текущего фрагмента."
    )

    add_black_heading(doc, "Алгоритмы", level=1)
    add_paragraph(doc, ", ".join(ALGORITHM_LABELS.values()) + ".")

    add_black_heading(doc, "Результаты", level=1)
    add_summary_table(doc, analysis)

    add_black_heading(doc, "Краткий анализ", level=1)
    summary = analysis["summary"]
    for array_type in ARRAY_LABELS:
        item = summary[array_type]
        fastest = item["fastest_3000"]
        least_chars = item["least_chars_3000"]
        add_paragraph(
            doc,
            f"Для типа «{ARRAY_LABELS[array_type]}» при n=3000 минимальное среднее время показал "
            f"{ALGORITHM_LABELS[fastest['algorithm']]}: "
            f"{format_num(float(fastest['mean_time_us']) / 1000.0, 2)} мс. "
            f"Минимальное число символьных операций у "
            f"{ALGORITHM_LABELS[least_chars['algorithm']]}: "
            f"{format_num(float(least_chars['mean_char_comparisons']) / 1000.0, 1)} тыс."
        )
    add_paragraph(
        doc,
        "По CSV видно, что адаптированные строковые алгоритмы сокращают работу с символами: "
        "String Quick Sort быстро обрабатывает разбиения по текущему символу, а MSD Radix Sort "
        "даёт минимальное число символьных операций."
    )
    add_paragraph(
        doc,
        "Обычные Quick Sort и Merge Sort повторно сравнивают совпадающие префиксы строк. "
        "Гибрид MSD Radix + Quick снижает накладные расходы на малых фрагментах, где размер "
        "подмассива меньше 74."
    )

    add_graphs(doc)

    add_black_heading(doc, "Codeforces", level=1)
    add_ids_table(doc, ids)

    add_black_heading(doc, "Выводы", level=1)
    add_paragraph(
        doc,
        f"Файлы с исходными замерами: {RAW.name}, {AVERAGED.name}. Реализации StringGenerator "
        "и StringSortTester находятся в исходниках проекта."
    )
    add_black_heading(doc, "Репозиторий", level=1)
    link_paragraph = doc.add_paragraph("Ссылка на публичный репозиторий: ")
    link_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in link_paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    add_hyperlink(link_paragraph, REPOSITORY_URL, REPOSITORY_URL)

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT)
    print(f"wrote {REPORT}")


if __name__ == "__main__":
    build_report()
